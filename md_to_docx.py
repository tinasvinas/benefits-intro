# -*- coding: utf-8 -*-
"""将 WorkBuddy实训系统需求文档.md 转换为格式化的 DOCX 文档。"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = r"D:\工作\产品设计\2025码聊\vibe coding\WorkBuddy实训系统需求文档.md"
OUT = r"D:\工作\产品设计\2025码聊\vibe coding\WorkBuddy实训系统需求文档.docx"

doc = Document()

# 默认中文字体
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

def set_cn(run, font="Microsoft YaHei"):
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)

def add_runs(paragraph, text):
    """解析 **加粗** 片段并添加到段落。"""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            run = paragraph.add_run(p[2:-2])
            run.bold = True
            set_cn(run)
        else:
            run = paragraph.add_run(p)
            set_cn(run)

def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")

def parse_table(rows):
    """rows: 包含表头、分隔行、数据行的原始行列表。返回 (header, data)。"""
    # 去掉分隔行（第二行）
    content = [r for r in rows if not re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", r)]
    table_data = []
    for r in content:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        table_data.append(cells)
    if not table_data:
        return [], []
    return table_data[0], table_data[1:]

lines = []
with open(SRC, "r", encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
n = len(lines)
while i < n:
    line = lines[i]
    stripped = line.strip()

    # 跳过空行
    if stripped == "":
        i += 1
        continue

    # 跳过纯分隔线
    if re.match(r"^---+$", stripped):
        i += 1
        continue

    # 表格
    if is_table_row(line):
        tbl_rows = []
        while i < n and is_table_row(lines[i]):
            tbl_rows.append(lines[i])
            i += 1
        header, data = parse_table(tbl_rows)
        if header:
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            for j, h in enumerate(header):
                hdr[j].text = ""
                add_runs(hdr[j].paragraphs[0], h)
                for r in hdr[j].paragraphs[0].runs:
                    r.bold = True
            for row in data:
                cells = table.add_row().cells
                for j, c in enumerate(row):
                    cells[j].text = ""
                    add_runs(cells[j].paragraphs[0], c)
        continue

    # 标题
    m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        if level == 1:
            p = doc.add_heading(level=0)
            add_runs(p, text)
        elif level == 2:
            p = doc.add_heading(level=1)
            add_runs(p, text)
        elif level == 3:
            p = doc.add_heading(level=2)
            add_runs(p, text)
        else:
            p = doc.add_heading(level=3)
            add_runs(p, text)
        i += 1
        continue

    # 无序列表
    if re.match(r"^-\s+", stripped):
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, re.sub(r"^-\s+", "", stripped))
        i += 1
        continue

    # 有序列表
    if re.match(r"^\d+\.\s+", stripped):
        p = doc.add_paragraph(style="List Number")
        add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
        i += 1
        continue

    # 普通段落
    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

doc.save(OUT)
print("DOCX 已生成:", OUT)
