from __future__ import annotations

import re
from collections import Counter
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "WorkBuddy实训系统需求文档_模板排版确认稿.docx"


def heading_level(style_name: str) -> int | None:
    match = re.search(r"Heading\s+(\d+)$", style_name)
    return int(match.group(1)) if match else None


def main() -> None:
    document = Document(DOCX)
    headings: list[tuple[int, int, str, str, str]] = []
    for index, paragraph in enumerate(document.paragraphs):
        level = heading_level(paragraph.style.name)
        if level is None:
            continue
        num_id = "-"
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            num_id = str(paragraph._p.pPr.numPr.numId.val)
        headings.append((index, level, paragraph.style.name, num_id, paragraph.text.strip()))

    print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} headings={len(headings)}")
    print("HEADINGS")
    for item in headings:
        print("\t".join(map(str, item)))

    print("TERMS")
    terms = ["创建", "新建", "添加", "新增", "页面入口", "数据模型", "UI", "性能"]
    counter = Counter()
    for paragraph in document.paragraphs:
        for term in terms:
            counter[term] += paragraph.text.count(term)
    for term in terms:
        print(f"{term}\t{counter[term]}")

    print("TERM_SAMPLES")
    for index, paragraph in enumerate(document.paragraphs):
        if any(term in paragraph.text for term in ["创建", "新建", "添加", "页面入口", "数据模型", "UI", "性能"]):
            print(f"{index}\t{paragraph.text.strip()}")

    print("LEVEL_JUMPS")
    previous_level = None
    for index, level, style, num_id, text in headings:
        if previous_level is not None and level > previous_level + 1:
            print(f"{index}\t{previous_level}->{level}\t{text}")
        previous_level = level


if __name__ == "__main__":
    main()
