from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
BEFORE = Path(__file__).resolve().parent / "WorkBuddy实训系统需求文档_模板排版确认稿.before-title-term-fix.docx"
AFTER = ROOT / "WorkBuddy实训系统需求文档_模板排版确认稿.docx"

NUMBER_RE = re.compile(r"^\s*\d+(?:\.\d+)*(?:\.)?\s+")


def normalized_expected(text: str, is_heading: bool) -> str:
    if is_heading:
        text = NUMBER_RE.sub("", text, count=1)
    text = text.replace("不创建或启动虚拟桌面", "不启动虚拟桌面")
    return text.replace("创建", "新增")


def main() -> None:
    before = Document(BEFORE)
    after = Document(AFTER)
    assert len(before.paragraphs) == len(after.paragraphs)
    assert len(before.tables) == len(after.tables)

    unexpected = []
    for index, (old, new) in enumerate(zip(before.paragraphs, after.paragraphs)):
        is_heading = "Heading" in old.style.name
        expected = normalized_expected(old.text, is_heading)
        if expected != new.text:
            unexpected.append((index, old.text, new.text, expected))

    before_cells = [cell for table in before.tables for row in table.rows for cell in row.cells]
    after_cells = [cell for table in after.tables for row in table.rows for cell in row.cells]
    assert len(before_cells) == len(after_cells)
    for index, (old_cell, new_cell) in enumerate(zip(before_cells, after_cells)):
        expected = normalized_expected(old_cell.text, False)
        if expected != new_cell.text:
            unexpected.append((f"table-cell-{index}", old_cell.text, new_cell.text, expected))

    print(f"paragraphs={len(after.paragraphs)} tables={len(after.tables)}")
    print(f"unexpected_text_changes={len(unexpected)}")
    for item in unexpected[:20]:
        print(item)
    if unexpected:
        raise SystemExit(2)


if __name__ == "__main__":
    main()
