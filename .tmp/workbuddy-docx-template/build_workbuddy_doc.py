from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "WorkBuddy实训系统需求文档.md"
TEMPLATE = Path(__file__).resolve().parent / "template.docx"
OUTPUT = ROOT / "WorkBuddy实训系统需求文档_模板排版确认稿.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TABLE_WIDTH_DXA = 8310


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = text.replace("\\`", "`")
    return text


def add_inline(paragraph, text: str, *, font_size: float | None = None):
    text = clean_inline(text)
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            if font_size:
                run.font.size = Pt(font_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
            run.font.size = Pt(font_size or 9.5)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        if font_size:
            run.font.size = Pt(font_size)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        if font_size:
            run.font.size = Pt(font_size)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]):
    widths = [max(700, int(x)) for x in widths]
    scale = TABLE_WIDTH_DXA / sum(widths)
    widths = [int(round(x * scale)) for x in widths]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_grid = tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        row.height = None
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_list_numbering(paragraph, num_id: int, level: int = 0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def new_decimal_numbering(doc) -> int:
    numbering = doc.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), "5")
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def copy_style_value(target, source, name):
    value = getattr(source, name)
    if value is not None:
        setattr(target, name, value)


def ensure_heading_styles(doc):
    for level in range(1, 6):
        source = doc.styles[f"Heading {level}"]
        style_name = f"WB Heading {level}"
        if style_name in doc.styles:
            continue
        target = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        target.base_style = doc.styles["Normal"]
        for attr in ("name", "size", "bold", "italic", "underline", "all_caps", "small_caps"):
            if attr == "name":
                if source.font.name:
                    target.font.name = source.font.name
                    target._element.rPr.rFonts.set(qn("w:eastAsia"), source.font.name)
            else:
                copy_style_value(target.font, source.font, attr)
        if source.font.color and source.font.color.type and source.font.color.rgb:
            target.font.color.rgb = source.font.color.rgb
        for attr in (
            "alignment",
            "first_line_indent",
            "keep_together",
            "keep_with_next",
            "left_indent",
            "line_spacing",
            "page_break_before",
            "right_indent",
            "space_after",
            "space_before",
            "widow_control",
        ):
            copy_style_value(target.paragraph_format, source.paragraph_format, attr)


def add_code_block(doc, lines: list[str]):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F8")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(9.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def set_body_paragraph(paragraph, *, after=Pt(5), before=Pt(0), line=1.15):
    paragraph.paragraph_format.space_before = before
    paragraph.paragraph_format.space_after = after
    paragraph.paragraph_format.line_spacing = line


def parse_table(lines: list[str], start: int):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def add_markdown_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=cols)
    table.style = "Table Grid"
    weights = []
    for column in range(cols):
        max_len = max(len(row[column]) for row in normalized)
        weights.append(max(10, min(max_len, 48)))
    if cols == 2:
        weights = [max(weights[0], 16), max(weights[1], 30)]
    elif cols == 3:
        weights = [14, 20, 48]
    set_table_geometry(table, weights)
    mark_header_row(table.rows[0])
    for row_index, row in enumerate(normalized):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            add_inline(paragraph, value, font_size=9.5)
            if row_index == 0:
                set_cell_shading(cell, "EAF2F8")
                for run in paragraph.runs:
                    run.bold = True
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_metadata_table(doc, metadata: list[tuple[str, str]]):
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [20, 80])
    for row, (key, value) in zip(table.rows, metadata):
        row.cells[0].text = ""
        row.cells[1].text = ""
        p0 = row.cells[0].paragraphs[0]
        p1 = row.cells[1].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        add_inline(p0, key, font_size=9.5)
        add_inline(p1, value, font_size=9.5)
        set_cell_shading(row.cells[0], "F4F6F8")
        for run in p0.runs:
            run.bold = True


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def main():
    source_lines = SOURCE.read_text(encoding="utf-8-sig").splitlines()
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    ensure_heading_styles(doc)

    # Preserve the template's A4 geometry and adjust only the body content.
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("WorkBuddy 实训系统需求文档")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("产品需求文档 · 模板排版确认稿")

    metadata = []
    in_meta = False
    for line in source_lines:
        if line.strip() == "| 项目 | 内容 |":
            in_meta = True
            continue
        if in_meta and line.startswith("|------"):
            continue
        if in_meta and line.strip().startswith("|"):
            parts = [part.strip() for part in line.strip().strip("|").split("|")]
            if len(parts) == 2:
                metadata.append((parts[0], parts[1]))
            continue
        if in_meta and line.strip() == "---":
            break
    add_metadata_table(doc, metadata)
    doc.add_paragraph()

    body_start = next((i + 1 for i, line in enumerate(source_lines) if line.strip() == "---"), 0)
    index = body_start
    ordered_num_id = None
    while index < len(source_lines):
        line = source_lines[index]
        stripped = line.strip()
        # Blank lines inside Markdown lists are visual separators, not list
        # boundaries. Keep the active ordered-list numbering across them;
        # headings, prose, tables, code blocks and separators below still
        # reset it when they are actually encountered.
        if not stripped:
            index += 1
            continue
        if stripped == "---" or stripped.startswith("# WorkBuddy 实训系统需求文档"):
            ordered_num_id = None
            index += 1
            continue
        # The source file may carry a historical Demo revision log. It is
        # intentionally omitted from the template-formatted requirements
        # document per the standing documentation rule.
        if re.match(r"^##\s+11\.", stripped):
            break
        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(source_lines) and not source_lines[index].strip().startswith("```"):
                code_lines.append(source_lines[index])
                index += 1
            if index < len(source_lines):
                index += 1
            add_code_block(doc, code_lines)
            ordered_num_id = None
            continue
        if stripped.startswith("|") and index + 1 < len(source_lines) and "|" in source_lines[index + 1]:
            rows, next_index = parse_table(source_lines, index)
            if len(rows) >= 2:
                add_markdown_table(doc, rows)
                index = next_index
                continue
        heading = re.match(r"^(#{2,5})\s+(.+)$", stripped)
        if heading:
            level = min(5, len(heading.group(1)) - 1)
            paragraph = doc.add_paragraph(style=f"WB Heading {level}")
            paragraph.add_run(heading.group(2).strip())
            ordered_num_id = None
            index += 1
            continue
        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        ordered = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if bullet or ordered:
            match = bullet or ordered
            level = min(2, len(match.group(1)) // 2)
            paragraph = doc.add_paragraph(style="List Paragraph")
            if bullet:
                add_list_numbering(paragraph, 1, level)
                if level == 0:
                    ordered_num_id = None
            else:
                if ordered_num_id is None:
                    ordered_num_id = new_decimal_numbering(doc)
                add_list_numbering(paragraph, ordered_num_id, level)
            set_body_paragraph(paragraph, after=Pt(2), line=1.08)
            add_inline(paragraph, match.group(2), font_size=10.5)
            index += 1
            continue
        ordered_num_id = None
        paragraph = doc.add_paragraph(style="Normal")
        set_body_paragraph(paragraph)
        add_inline(paragraph, stripped, font_size=10.5)
        index += 1

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
