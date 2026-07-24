from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "WorkBuddy实训系统需求文档.md"
OUTPUT = ROOT / "WorkBuddy实训系统需求文档_模板排版确认稿.docx"


def plain_markdown(line: str) -> str:
    text = line.strip()
    text = re.sub(r"^#{1,6}\s+", "", text)
    text = re.sub(r"^\d+[.)]\s+", "", text)
    text = re.sub(r"^[-*]\s+", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("`", "")
    return text


def main() -> None:
    document = Document(OUTPUT)
    paragraphs = [p.text for p in document.paragraphs]
    cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    full_text = "\n".join([*paragraphs, *cells])

    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    content_lines: list[str] = []
    in_revision_log = False
    in_code = False
    for line in source_lines:
        stripped = line.strip()
        if re.match(r"^##\s+11\.", stripped):
            in_revision_log = True
        if in_revision_log:
            continue
        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if not stripped or stripped == "---" or stripped == "**文档结束**":
            continue
        if stripped.startswith("# WorkBuddy 实训系统需求文档"):
            continue
        if stripped.startswith("|") and set(stripped) <= {"|", "-", ":", " "}:
            continue
        if stripped.startswith("|"):
            parts = [plain_markdown(part.strip()) for part in stripped.strip("|").split("|")]
            content_lines.extend(part for part in parts if part)
            continue
        content_lines.append(plain_markdown(line))

    # Code blocks are intentionally converted to styled multi-line paragraphs,
    # so compare their individual paths separately after normalizing slashes.
    missing = [line for line in content_lines if len(line) >= 4 and line not in full_text]
    missing = [line for line in missing if line.rstrip("\\") not in full_text]
    print(f"paragraphs={len(paragraphs)} tables={len(document.tables)}")
    print(f"source_items={len(content_lines)} missing={len(missing)}")
    for item in missing[:30]:
        print(f"MISSING: {item}")
    print(f"revision_log_present={'2026-07-23 Demo 修订说明' in full_text}")
    print(f"document_end_present={'文档结束' in full_text}")

    targets = {
        "项目概览区": "实训指标区",
        "左侧任务导航": "顶部项目状态栏",
        "Skills": "MCP",
    }
    for first, second in targets.items():
        p1 = next(p for p in document.paragraphs if p.text == first)
        p2 = next(p for p in document.paragraphs if p.text == second)
        n1 = p1._p.pPr.numPr.numId.val
        n2 = p2._p.pPr.numPr.numId.val
        print(f"numbering_group {first}/{second}: {n1}/{n2} same={n1 == n2}")

    if missing:
        raise SystemExit(2)
    if "2026-07-23 Demo 修订说明" in full_text or "文档结束" in full_text:
        raise SystemExit(3)


if __name__ == "__main__":
    main()
