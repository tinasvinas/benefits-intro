from __future__ import annotations

import re
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "WorkBuddy实训系统需求文档_模板排版确认稿.docx"

HEADING_STYLE_RE = re.compile(r"(?:WB\s+)?Heading\s+(\d+)$")
MANUAL_NUMBER_RE = re.compile(r"^\s*\d+(?:\.\d+)*(?:\.)?\s+")


def replace_text_in_paragraph(paragraph, old: str, new: str) -> int:
    replacements = 0
    for run in paragraph.runs:
        count = run.text.count(old)
        if count:
            run.text = run.text.replace(old, new)
            replacements += count
    return replacements


def iter_all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def main() -> None:
    document = Document(DOCX)
    heading_changes = 0
    number_prefix_changes = 0

    for paragraph in document.paragraphs:
        match = HEADING_STYLE_RE.search(paragraph.style.name)
        if not match:
            continue
        level = int(match.group(1))
        target_style = f"Heading {level}"
        if paragraph.style.name != target_style:
            paragraph.style = target_style
            heading_changes += 1

        cleaned = MANUAL_NUMBER_RE.sub("", paragraph.text, count=1)
        if cleaned != paragraph.text:
            paragraph.text = cleaned
            paragraph.style = target_style
            number_prefix_changes += 1

    wording_changes = 0
    for paragraph in iter_all_paragraphs(document):
        if "不新增或启动虚拟桌面" in paragraph.text:
            paragraph.text = paragraph.text.replace("不新增或启动虚拟桌面", "不启动虚拟桌面")
            wording_changes += 1
        # This sentence describes runtime behavior, not the creation of a new
        # business object, so remove the ambiguous verb instead of replacing it.
        wording_changes += replace_text_in_paragraph(
            paragraph,
            "不创建或启动虚拟桌面",
            "不启动虚拟桌面",
        )
        wording_changes += replace_text_in_paragraph(
            paragraph,
            "不新增或启动虚拟桌面",
            "不启动虚拟桌面",
        )
        wording_changes += replace_text_in_paragraph(paragraph, "创建", "新增")

    document.save(DOCX)
    print(f"heading_style_changes={heading_changes}")
    print(f"manual_number_prefix_changes={number_prefix_changes}")
    print(f"wording_changes={wording_changes}")
    print(DOCX)


if __name__ == "__main__":
    main()
